"""
Ingestion tests: parser and chunker. Indexer/embedder require model and ChromaDB.
"""
from __future__ import annotations

import pytest

from app.services.ingestion.parser import parse_document, parse_csv, parse_docx
from app.services.ingestion.chunker import chunk_pages, MIN_CHUNK_TOKENS


def test_parse_csv() -> None:
    data = b"col1,col2\na,b\nc,d"
    pages = parse_csv(data)
    assert len(pages) == 1
    assert pages[0]["page_number"] == 1
    assert "a\tb" in pages[0]["text"]


def test_parse_docx() -> None:
    from docx import Document
    from io import BytesIO
    doc = Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("Second para.")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    pages = parse_docx(buf.read())
    assert len(pages) == 1
    assert "Hello world" in pages[0]["text"]
    assert "Second para" in pages[0]["text"]


def test_parse_document_unsupported() -> None:
    with pytest.raises(ValueError, match="Unsupported"):
        parse_document(b"x", "unknown")


def test_chunk_pages_empty() -> None:
    assert chunk_pages([]) == []
    assert chunk_pages([{"page_number": 1, "text": ""}]) == []
    assert chunk_pages([{"page_number": 1, "text": "   \n  "}]) == []


def test_chunk_pages_small() -> None:
    # One short paragraph -> one chunk if above MIN
    text = " ".join(["word"] * 50)
    pages = [{"page_number": 1, "text": text}]
    chunks = chunk_pages(pages)
    assert len(chunks) >= 1
    assert chunks[0].page_number == 1
    assert chunks[0].chunk_index == 0
    assert chunks[0].token_count >= MIN_CHUNK_TOKENS or len(chunks) == 1


def test_chunk_pages_large() -> None:
    # Many words -> multiple chunks with overlap
    text = " ".join(["word"] * 800)
    pages = [{"page_number": 1, "text": text}]
    chunks = chunk_pages(pages)
    assert len(chunks) >= 2
    assert all(c.page_number == 1 for c in chunks)
    assert [c.chunk_index for c in chunks] == list(range(len(chunks)))
